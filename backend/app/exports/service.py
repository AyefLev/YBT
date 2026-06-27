from datetime import datetime, timezone
from html import unescape
from pathlib import Path
import re
from uuid import uuid4

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.orm import Session

from app.ai.formatting import normalize_generated_math_text
from app.auth.models import User
from app.core.config import get_settings
from app.courses.models import Course
from app.exports.models import ExportRecord
from app.exercises.models import Exercise
from app.lessons.models import Lesson
from app.questions.models import QuestionBankItem
from app.questions.service import decode_list


def export_lesson_docx(
    db: Session,
    *,
    lesson: Lesson,
    current_user: User,
) -> Path:
    metadata = [
        ("标题", lesson.title),
        ("学科", lesson.subject),
        ("章节", lesson.chapter),
        ("阶段", lesson.stage),
        ("课时", f"{lesson.duration_minutes} 分钟"),
        ("学情", lesson.student_level),
        ("合规风险", lesson.compliance_level),
    ]
    return _write_docx_and_record(
        db,
        resource_type="lesson",
        resource_id=lesson.id,
        title=lesson.title,
        metadata=metadata,
        content=lesson.current_content,
        current_user=current_user,
    )


def export_exercise_docx(
    db: Session,
    *,
    exercise: Exercise,
    current_user: User,
) -> Path:
    metadata = [
        ("标题", exercise.title),
        ("学科", exercise.subject),
        ("知识点", exercise.knowledge_point),
        ("题型", exercise.question_type),
        ("难度", exercise.difficulty),
        ("合规风险", exercise.compliance_level),
    ]
    return _write_docx_and_record(
        db,
        resource_type="exercise",
        resource_id=exercise.id,
        title=exercise.title,
        metadata=metadata,
        content=exercise.current_content,
        current_user=current_user,
    )


def export_course_outline_docx(
    db: Session,
    *,
    course: Course,
    current_user: User,
) -> Path:
    document = Document()
    document.add_heading(course.title, level=1)
    document.add_paragraph(f"学科：{course.subject}")
    document.add_paragraph(f"考试类型：{course.exam_type}")
    document.add_paragraph(f"状态：{course.status}")
    if course.description:
        document.add_paragraph(f"课程说明：{course.description}")

    document.add_heading("章节与课次", level=2)
    if not course.chapters:
        document.add_paragraph("暂无章节。")
    for chapter in course.chapters:
        document.add_heading(f"{chapter.order_index}. {chapter.title}", level=3)
        if chapter.summary:
            document.add_paragraph(chapter.summary)
        for session in chapter.sessions:
            document.add_paragraph(
                f"{session.order_index}. {session.title}（{session.duration_minutes} 分钟）"
            )
            if session.teaching_goal:
                document.add_paragraph(f"教学目标：{session.teaching_goal}")

    document.add_heading("知识点", level=2)
    if not course.knowledge_points:
        document.add_paragraph("暂无知识点。")
    for point in course.knowledge_points:
        suffix = point.description or point.difficulty
        document.add_paragraph(f"{point.name}：{suffix}")

    return _save_document_and_record(
        db,
        document=document,
        resource_type="course",
        resource_id=course.id,
        title=course.title,
        current_user=current_user,
    )


def export_question_package_docx(
    db: Session,
    *,
    questions: list[QuestionBankItem],
    current_user: User,
) -> Path:
    document = Document()
    document.add_heading("题库练习包", level=1)
    for index, question in enumerate(questions, start=1):
        document.add_heading(f"{index}. {question.title}", level=2)
        document.add_paragraph(f"学科：{question.subject}")
        document.add_paragraph(f"题型：{question.question_type}")
        document.add_paragraph(f"难度：{question.difficulty}")
        document.add_paragraph(question.stem)
        for option in decode_list(question.options_json):
            document.add_paragraph(option)
        if question.answer:
            document.add_paragraph(f"答案：{question.answer}")
        if question.analysis:
            document.add_paragraph(f"解析：{question.analysis}")

    return _save_document_and_record(
        db,
        document=document,
        resource_type="question_package",
        resource_id=0,
        title="question-package",
        current_user=current_user,
    )


def _write_docx_and_record(
    db: Session,
    *,
    resource_type: str,
    resource_id: int,
    title: str,
    metadata: list[tuple[str, str]],
    content: str,
    current_user: User,
) -> Path:
    export_dir = get_settings().export_dir
    export_dir.mkdir(parents=True, exist_ok=True)
    file_path = _export_file_path(export_dir, resource_type, resource_id, title)

    try:
        document = Document()
        document.add_heading(title, level=1)
        for label, value in metadata:
            document.add_paragraph(f"{label}：{value}")
        document.add_heading("正文", level=2)
        _add_content(document, content)
        document.save(file_path)

        db.add(
            ExportRecord(
                resource_type=resource_type,
                resource_id=resource_id,
                file_path=str(file_path),
                exported_by=current_user.id,
            )
        )
        db.flush()
    except Exception:
        remove_export_file(file_path)
        raise
    return file_path


def _save_document_and_record(
    db: Session,
    *,
    document: Document,
    resource_type: str,
    resource_id: int,
    title: str,
    current_user: User,
) -> Path:
    export_dir = get_settings().export_dir
    export_dir.mkdir(parents=True, exist_ok=True)
    file_path = _export_file_path(export_dir, resource_type, resource_id, title)

    try:
        document.save(file_path)
        db.add(
            ExportRecord(
                resource_type=resource_type,
                resource_id=resource_id,
                file_path=str(file_path),
                exported_by=current_user.id,
            )
        )
        db.flush()
    except Exception:
        remove_export_file(file_path)
        raise
    return file_path


def _export_file_path(
    export_dir: Path,
    resource_type: str,
    resource_id: int,
    title: str,
) -> Path:
    safe_title = re.sub(r"[^A-Za-z0-9._-]+", "-", title).strip("-") or resource_type
    for _ in range(10):
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S%f")
        candidate = (
            export_dir
            / f"{resource_type}-{resource_id}-{safe_title}-{timestamp}-{uuid4().hex}.docx"
        )
        if not candidate.exists():
            return candidate
    raise RuntimeError("Could not allocate a unique export filename")


def remove_export_file(file_path: Path) -> None:
    try:
        file_path.unlink(missing_ok=True)
    except OSError:
        pass


def _add_content(document: Document, content: str) -> None:
    if not content.strip():
        return

    pattern = re.compile(r"\[(公式|鍏紡|SVG)\]([\s\S]*?)\[/\1\]", re.IGNORECASE)
    last_index = 0
    for match in pattern.finditer(content):
        _add_text_content(document, content[last_index : match.start()])
        block_type = match.group(1).upper()
        block_body = match.group(2).strip()
        if block_type == "SVG":
            _add_svg_summary(document, block_body)
        else:
            _add_formula_block(document, block_body)
        last_index = match.end()

    _add_text_content(document, content[last_index:])


def _add_text_content(document: Document, content: str) -> None:
    for raw_line in content.splitlines() or [content]:
        line = _clean_markdown_line(raw_line)
        if not line:
            continue
        if _is_heading(line):
            document.add_heading(line, level=3)
        else:
            document.add_paragraph(line)


def _add_formula_block(document: Document, formula: str) -> None:
    matrices = _extract_latex_matrices(formula)
    if matrices and _formula_without_matrices(formula).strip() == "":
        for matrix in matrices:
            _add_matrix_text(document, matrix)
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(_latex_to_readable(formula))
    run.italic = True


def _add_matrix_text(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    width = max(len(row) for row in rows)
    normalized_rows = [
        [_latex_to_readable(row[index]) if index < len(row) else "" for index in range(width)]
        for row in rows
    ]
    column_widths = [
        max(len(row[index]) for row in normalized_rows)
        for index in range(width)
    ]

    for row_index, row in enumerate(normalized_rows):
        left = "⎡" if row_index == 0 else "⎣" if row_index == len(normalized_rows) - 1 else "⎢"
        right = "⎤" if row_index == 0 else "⎦" if row_index == len(normalized_rows) - 1 else "⎥"
        body = "  ".join(
            value.center(column_widths[index])
            for index, value in enumerate(row)
        )
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(f"{left} {body} {right}")
        run.font.name = "Consolas"


def _add_svg_summary(document: Document, svg: str) -> None:
    labels = [
        _clean_markdown_line(unescape(match.strip()))
        for match in re.findall(r"<text\b[^>]*>([\s\S]*?)</text>", svg, flags=re.IGNORECASE)
        if match.strip()
    ]
    description = "、".join(label for label in labels if label)
    if description:
        document.add_paragraph(f"【示意图】{description}")
    else:
        document.add_paragraph("【示意图】请在系统预览中查看图形内容。")


def _extract_latex_matrices(formula: str) -> list[list[list[str]]]:
    matrices: list[list[list[str]]] = []
    pattern = re.compile(r"\\begin\{[bpv]?matrix\}([\s\S]*?)\\end\{[bpv]?matrix\}")
    for match in pattern.finditer(formula):
        rows = [
            [_latex_to_readable(cell.strip()) for cell in row.split("&")]
            for row in re.split(r"\\\\", match.group(1))
            if row.strip()
        ]
        if rows:
            matrices.append(rows)
    return matrices


def _formula_without_matrices(formula: str) -> str:
    return re.sub(r"\\begin\{[bpv]?matrix\}[\s\S]*?\\end\{[bpv]?matrix\}", "", formula)


def _latex_to_readable(value: str) -> str:
    value = value.strip()
    replacements = {
        r"\times": "×",
        r"\cdot": "·",
        r"\div": "÷",
        r"\leq": "≤",
        r"\geq": "≥",
        r"\neq": "≠",
        r"\pm": "±",
        r"\left": "",
        r"\right": "",
    }
    for source, target in replacements.items():
        value = value.replace(source, target)
    value = value.replace(r"\[", "").replace(r"\]", "")
    value = value.replace(r"\(", "").replace(r"\)", "")
    value = re.sub(r"\\begin\{[bpv]?matrix\}", "[", value)
    value = re.sub(r"\\end\{[bpv]?matrix\}", "]", value)
    value = value.replace("\\\\", "; ")
    value = value.replace("&", "  ")
    value = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", value)
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def _clean_markdown_line(line: str) -> str:
    line = normalize_generated_math_text(line).strip()
    if not line or set(line) <= {"-"}:
        return ""
    line = re.sub(r"^#{1,6}\s*", "", line)
    line = re.sub(r"^\*\s+", "", line)
    line = line.replace("**", "")
    line = line.replace("[公式]", "").replace("[/公式]", "")
    line = line.replace("[鍏紡]", "").replace("[/鍏紡]", "")
    line = line.replace("[SVG]", "").replace("[/SVG]", "")
    line = line.replace(r"\[", "").replace(r"\]", "")
    line = line.replace(r"\(", "").replace(r"\)", "")
    return line.strip()


def _is_heading(line: str) -> bool:
    return bool(
        re.match(
            r"^(课堂练习|课程设计|教学目标|重点难点|教学流程|课后建议|题目\s*\d+|第[一二三四五六七八九十]+题|答案与解析)",
            line,
        )
        or line.endswith("：")
    )
