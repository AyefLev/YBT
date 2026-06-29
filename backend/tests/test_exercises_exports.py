from pathlib import Path
from io import BytesIO

import pytest
from docx import Document


def _auth_headers(client, username: str = "teacher_exercises") -> dict[str, str]:
    client.post(
        "/api/auth/register",
        json={
            "username": username,
            "email": f"{username}@example.com",
            "password": "secret-password",
            "display_name": "Exercise Teacher",
        },
    )
    login_response = client.post(
        "/api/auth/login",
        json={"username": username, "password": "secret-password"},
    )
    token = login_response.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def _exercise_generate_payload(**overrides):
    payload = {
        "title": "Cloze practice",
        "subject": "English",
        "knowledge_point": "context clues",
        "question_type": "short answer",
        "difficulty": "medium",
        "count": 3,
        "use_materials": False,
        "material_ids": [],
    }
    payload.update(overrides)
    return payload


def _exercise_save_payload(**overrides):
    payload = {
        "title": "Cloze practice",
        "subject": "English",
        "knowledge_point": "context clues",
        "question_type": "short answer",
        "difficulty": "medium",
        "content": "1. Explain the context clue.\n答案: infer from nearby words.",
    }
    payload.update(overrides)
    return payload


def _lesson_save_payload(**overrides):
    payload = {
        "title": "Lesson export",
        "subject": "Writing",
        "chapter": "Paragraphs",
        "stage": "Graduate exam",
        "duration_minutes": 45,
        "student_level": "Basic",
        "content": "Teaching content for export.",
    }
    payload.update(overrides)
    return payload


def _upload_txt(client, headers: dict[str, str], title: str, content: str):
    return client.post(
        "/api/materials/upload",
        headers=headers,
        data={"title": title},
        files={"file": ("exercise.txt", content.encode("utf-8"), "text/plain")},
    )


def _create_exercise(client, headers: dict[str, str], **overrides):
    response = client.post(
        "/api/exercises",
        headers=headers,
        json=_exercise_save_payload(**overrides),
    )
    assert response.status_code == 201
    return response.json()


def _create_lesson(client, headers: dict[str, str], **overrides):
    response = client.post(
        "/api/lessons",
        headers=headers,
        json=_lesson_save_payload(**overrides),
    )
    assert response.status_code == 201
    return response.json()


def _create_question(client, headers: dict[str, str], **overrides):
    payload = {
        "title": "极限练习",
        "subject": "数学",
        "question_type": "single_choice",
        "difficulty": "basic",
        "stem": r"已知函数 \( f(x)=\frac{\int_{0}^{x}|\sin t|\,dt}{x^{\alpha}} \)，求 \(\alpha\) 的范围。",
        "options": [r"A. \( [0,2] \)", r"B. \( [1,2] \)", r"C. \( [1,3] \)", r"D. \( [0,3] \)"],
        "answer": "B",
        "analysis": r"当 \(x \to 0\) 时，分子约为 \(x^2/2\)，因此需要 \(\alpha \leq 2\)。",
        "tags": ["极限"],
    }
    payload.update(overrides)
    response = client.post("/api/questions", headers=headers, json=payload)
    assert response.status_code == 201
    return response.json()


def test_generate_exercise_without_materials_returns_content_status_and_compliance(client):
    headers = _auth_headers(client)

    response = client.post(
        "/api/exercises/generate",
        headers=headers,
        json=_exercise_generate_payload(),
    )

    assert response.status_code == 200
    body = response.json()
    assert "答案" in body["content"]
    assert body["references"] == []
    assert body["provider_status"]["provider"] == "mock"
    assert body["compliance"]["risk_level"] == "low"


def test_save_exercise_creates_exercise_and_version_one(client):
    headers = _auth_headers(client, username="teacher_save_exercise")

    create_response = client.post(
        "/api/exercises",
        headers=headers,
        json=_exercise_save_payload(),
    )

    assert create_response.status_code == 201
    exercise = create_response.json()
    assert exercise["id"]
    assert exercise["title"] == "Cloze practice"
    assert exercise["current_content"].startswith("1. Explain")
    assert exercise["compliance_level"] == "low"

    versions_response = client.get(
        f"/api/exercises/{exercise['id']}/versions",
        headers=headers,
    )
    assert versions_response.status_code == 200
    versions = versions_response.json()
    assert len(versions) == 1
    assert versions[0]["version_no"] == 1
    assert versions[0]["content"] == exercise["current_content"]


def test_restore_exercise_version_updates_current_content_and_preserves_history(client):
    headers = _auth_headers(client, username="teacher_restore_exercise")
    exercise = _create_exercise(client, headers, content="Original content\n答案: A")
    versions = client.get(
        f"/api/exercises/{exercise['id']}/versions",
        headers=headers,
    ).json()

    response = client.post(
        f"/api/exercises/{exercise['id']}/restore-version",
        headers=headers,
        json={"version_id": versions[0]["id"]},
    )

    assert response.status_code == 200
    assert response.json()["current_content"] == "Original content\n答案: A"
    versions_response = client.get(
        f"/api/exercises/{exercise['id']}/versions",
        headers=headers,
    )
    restored_versions = versions_response.json()
    assert [version["version_no"] for version in restored_versions] == [1, 2]
    assert restored_versions[1]["content"] == "Original content\n答案: A"


def test_cross_user_cannot_access_another_users_exercise_or_versions(client):
    owner_headers = _auth_headers(client, username="teacher_exercise_owner")
    other_headers = _auth_headers(client, username="teacher_exercise_other")
    exercise = _create_exercise(client, owner_headers, title="Private Exercise")

    exercise_response = client.get(f"/api/exercises/{exercise['id']}", headers=other_headers)
    versions_response = client.get(
        f"/api/exercises/{exercise['id']}/versions",
        headers=other_headers,
    )

    assert exercise_response.status_code == 404
    assert versions_response.status_code == 404


def test_generate_exercise_with_materials_returns_owner_scoped_reference_metadata(client):
    owner_headers = _auth_headers(client, username="teacher_exercise_material")
    other_headers = _auth_headers(client, username="teacher_exercise_other_material")
    owner_material = _upload_txt(
        client,
        owner_headers,
        "Owner Exercise Material",
        "context clues require reading nearby signal words",
    ).json()
    other_material = _upload_txt(
        client,
        other_headers,
        "Other Exercise Material",
        "context clues private other user material",
    ).json()

    response = client.post(
        "/api/exercises/generate",
        headers=owner_headers,
        json=_exercise_generate_payload(
            knowledge_point="context clues",
            use_materials=True,
            material_ids=[owner_material["id"], other_material["id"]],
        ),
    )

    assert response.status_code == 200
    references = response.json()["references"]
    assert references
    assert any(
        reference["material_id"] == owner_material["id"]
        and "nearby signal words" in reference["content"]
        for reference in references
    )
    assert all("id" in reference for reference in references)
    assert all("page_no" in reference for reference in references)
    assert all("slide_no" in reference for reference in references)
    assert all("score" in reference for reference in references)
    assert all(reference["material_id"] != other_material["id"] for reference in references)


def test_phase2_generate_exercise_returns_multi_ai_review(client, monkeypatch):
    monkeypatch.setenv("LLM_MULTI_AGENT_REVIEW", "true")
    from app.core.config import get_settings

    get_settings.cache_clear()
    headers = _auth_headers(client, username="teacher_exercise_review")

    response = client.post(
        "/api/exercises/generate",
        headers=headers,
        json=_exercise_generate_payload(),
    )

    assert response.status_code == 200
    body = response.json()
    assert body["review"]["enabled"] is True
    assert body["review"]["status"] in {"passed", "warning"}
    assert body["review"]["reviewer_model"]
    assert "warnings" in body["review"]
    assert "suggestions" in body["review"]
    assert "raw_review" in body["review"]


def _assert_docx_attachment(response):
    assert response.status_code == 200
    assert response.content.startswith(b"PK")
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    content_disposition = response.headers["content-disposition"]
    assert "attachment" in content_disposition
    assert ".docx" in content_disposition


def _docx_text(response) -> str:
    document = Document(BytesIO(response.content))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def _docx_all_text(response) -> str:
    document = Document(BytesIO(response.content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join(paragraphs + table_cells)


def _docx_table_count(response) -> int:
    document = Document(BytesIO(response.content))
    return len(document.tables)


def test_export_saved_lesson_to_docx_returns_file_response_and_creates_record(client):
    from app.core.database import get_session_local
    from app.exports.models import ExportRecord

    headers = _auth_headers(client, username="teacher_export_lesson")
    lesson = _create_lesson(client, headers)

    response = client.post(
        f"/api/exports/lesson/{lesson['id']}/docx",
        headers=headers,
    )

    _assert_docx_attachment(response)

    session_local = get_session_local()
    with session_local() as db:
        records = db.query(ExportRecord).all()
        assert len(records) == 1
        assert records[0].resource_type == "lesson"
        assert records[0].resource_id == lesson["id"]
        assert Path(records[0].file_path).exists()


def test_export_saved_exercise_to_docx_returns_file_response_and_creates_record(client):
    from app.core.database import get_session_local
    from app.exports.models import ExportRecord

    headers = _auth_headers(client, username="teacher_export_exercise")
    exercise = _create_exercise(client, headers)

    response = client.post(
        f"/api/exports/exercise/{exercise['id']}/docx",
        headers=headers,
    )

    _assert_docx_attachment(response)

    session_local = get_session_local()
    with session_local() as db:
        records = db.query(ExportRecord).all()
        assert len(records) == 1
        assert records[0].resource_type == "exercise"
        assert records[0].resource_id == exercise["id"]
    assert Path(records[0].file_path).exists()


def test_export_docx_uses_chinese_labels_and_cleans_markdown(client):
    headers = _auth_headers(client, username="teacher_export_polished")
    exercise = _create_exercise(
        client,
        headers,
        title="课堂练习",
        subject="数学",
        knowledge_point="矩阵乘法",
        question_type="选择题",
        difficulty="基础",
        content=(
            "### **课堂练习：矩阵乘法（基础）**\n"
            "---\n"
            "#### **题目 1**\n"
            "*   **正确答案：A**\n"
            "解析：矩阵乘法需要维数相容。"
        ),
    )

    response = client.post(
        f"/api/exports/exercise/{exercise['id']}/docx",
        headers=headers,
    )

    _assert_docx_attachment(response)
    text = _docx_text(response)
    assert "标题：课堂练习" in text
    assert "学科：数学" in text
    assert "知识点：矩阵乘法" in text
    assert "正文" in text
    assert "Title:" not in text
    assert "Subject:" not in text
    assert "Content" not in text
    assert "###" not in text
    assert "**" not in text
    assert "---" not in text
    assert "课堂练习：矩阵乘法（基础）" in text
    assert "正确答案：A" in text


def test_export_docx_renders_formula_and_svg_protocol_readably(client):
    headers = _auth_headers(client, username="teacher_export_formula_protocol")
    exercise = _create_exercise(
        client,
        headers,
        title="矩阵乘法综合练习",
        subject="数学",
        knowledge_point="矩阵乘法",
        question_type="综合题",
        difficulty="中等",
        content=(
            "题目 1\n"
            "已知矩阵\n"
            "[公式]\\begin{pmatrix} 1 & 2 \\\\ 3 & 4 \\end{pmatrix}[/公式]\n"
            "请计算它与列向量相乘的结果。\n"
            "[SVG]<svg viewBox=\"0 0 100 40\"><text x=\"10\" y=\"20\">A 到 B</text></svg>[/SVG]\n"
            "答案与解析\n"
            "按矩阵乘法行乘列计算。"
        ),
    )

    response = client.post(
        f"/api/exports/exercise/{exercise['id']}/docx",
        headers=headers,
    )

    _assert_docx_attachment(response)
    text = _docx_all_text(response)
    assert "[公式]" not in text
    assert "[/公式]" not in text
    assert "[SVG]" not in text
    assert "[/SVG]" not in text
    assert "\\begin{pmatrix}" not in text
    assert "<svg" not in text
    assert "1" in text
    assert "2" in text
    assert "3" in text
    assert "4" in text
    assert "示意图" in text
    assert "A 到 B" in text
    assert _docx_table_count(response) == 0


def test_export_docx_cleans_inline_latex_math_commands(client):
    headers = _auth_headers(client, username="teacher_export_inline_latex")
    exercise = _create_exercise(
        client,
        headers,
        title="课堂练习",
        subject="数学",
        knowledge_point="矩阵乘法",
        question_type="证明题",
        difficulty="中等",
        content=(
            r"1. 设 \(A\) 为 \(m \times n\) 矩阵，\(B\) 为 \(n \times p\) 矩阵。"
            r"证明 \((AB)^{\mathrm{T}} = B^{\mathrm{T}}A^{\mathrm{T}}\)。"
            "\n"
            r"2. 定义 \(\operatorname{tr}(X)=\sum_{i=1}^{p}x_{ii}\)，证明迹的性质。"
        ),
    )

    response = client.post(
        f"/api/exports/exercise/{exercise['id']}/docx",
        headers=headers,
    )

    _assert_docx_attachment(response)
    text = _docx_all_text(response)
    assert r"\(" not in text
    assert r"\times" not in text
    assert r"\mathrm" not in text
    assert r"\operatorname" not in text
    assert r"\sum" not in text
    assert "A 为 m × n 矩阵" in text
    assert "(AB)^T = B^TA^T" in text
    assert "tr(X)=Σ_i=1^px_ii" in text


def test_export_question_package_cleans_complex_math(client):
    headers = _auth_headers(client, username="teacher_export_question_math")
    question = _create_question(client, headers)

    response = client.post(
        "/api/exports/questions/docx",
        headers=headers,
        json={"question_ids": [question["id"]]},
    )

    _assert_docx_attachment(response)
    text = _docx_all_text(response)
    assert r"\frac" not in text
    assert r"\int" not in text
    assert r"\alpha" not in text
    assert r"\leq" not in text
    assert "∫_0^x|sin t| dt" in text
    assert "α" in text
    assert "≤" in text


def test_exports_use_distinct_uuid_filenames(client):
    from app.core.database import get_session_local
    from app.exports.models import ExportRecord

    headers = _auth_headers(client, username="teacher_export_uuid")
    lesson = _create_lesson(client, headers, title="Repeat Export")

    first_response = client.post(
        f"/api/exports/lesson/{lesson['id']}/docx",
        headers=headers,
    )
    second_response = client.post(
        f"/api/exports/lesson/{lesson['id']}/docx",
        headers=headers,
    )

    _assert_docx_attachment(first_response)
    _assert_docx_attachment(second_response)

    session_local = get_session_local()
    with session_local() as db:
        paths = [Path(record.file_path) for record in db.query(ExportRecord).all()]
        assert len(paths) == 2
        assert paths[0] != paths[1]
        assert all(path.exists() for path in paths)


def test_cross_user_cannot_export_another_users_lesson_or_exercise(client):
    owner_headers = _auth_headers(client, username="teacher_export_owner")
    other_headers = _auth_headers(client, username="teacher_export_other")
    lesson = _create_lesson(client, owner_headers, title="Private Lesson")
    exercise = _create_exercise(client, owner_headers, title="Private Exercise")

    lesson_response = client.post(
        f"/api/exports/lesson/{lesson['id']}/docx",
        headers=other_headers,
    )
    exercise_response = client.post(
        f"/api/exports/exercise/{exercise['id']}/docx",
        headers=other_headers,
    )

    assert lesson_response.status_code == 404
    assert exercise_response.status_code == 404


def test_restore_exercise_version_refreshes_compliance_level(client):
    from app.compliance.service import HIGH_RISK_TERMS
    from app.core.database import get_session_local
    from app.exercises.models import Exercise
    from app.exercises.service import add_exercise_version

    headers = _auth_headers(client, username="teacher_restore_exercise_compliance")
    exercise = _create_exercise(client, headers, content="Safe exercise content")
    risky_content = f"Risky exercise content {HIGH_RISK_TERMS[0]}"

    session_local = get_session_local()
    with session_local() as db:
        version = add_exercise_version(
            db,
            exercise_id=exercise["id"],
            content=risky_content,
            change_note="risky",
        )
        version_id = version.id
        db.commit()

    response = client.post(
        f"/api/exercises/{exercise['id']}/restore-version",
        headers=headers,
        json={"version_id": version_id},
    )

    assert response.status_code == 200
    body = response.json()
    assert body["current_content"] == risky_content
    assert body["compliance_level"] == "high"
    with session_local() as db:
        restored = db.get(Exercise, exercise["id"])
        assert restored.compliance_level == "high"


def test_restore_lesson_version_refreshes_compliance_level(client):
    from app.compliance.service import HIGH_RISK_TERMS
    from app.core.database import get_session_local
    from app.lessons.models import Lesson
    from app.lessons.service import add_lesson_version

    headers = _auth_headers(client, username="teacher_restore_lesson_compliance")
    lesson = _create_lesson(client, headers, content="Safe lesson content")
    risky_content = f"Risky lesson content {HIGH_RISK_TERMS[0]}"

    session_local = get_session_local()
    with session_local() as db:
        version = add_lesson_version(
            db,
            lesson_id=lesson["id"],
            content=risky_content,
            change_note="risky",
        )
        version_id = version.id
        db.commit()

    response = client.post(
        f"/api/lessons/{lesson['id']}/restore-version",
        headers=headers,
        json={"version_id": version_id},
    )

    assert response.status_code == 200
    body = response.json()
    assert body["current_content"] == risky_content
    assert body["compliance_level"] == "high"
    with session_local() as db:
        restored = db.get(Lesson, lesson["id"])
        assert restored.compliance_level == "high"


def test_export_file_is_removed_when_record_flush_fails(client):
    from types import SimpleNamespace

    from app.core.config import get_settings
    from app.exports.service import export_lesson_docx

    class FailingDb:
        def add(self, value):
            self.value = value

        def flush(self):
            raise RuntimeError("record flush failed")

    lesson = SimpleNamespace(
        id=123,
        title="Cleanup Lesson",
        subject="Writing",
        chapter="Cleanup",
        stage="Graduate exam",
        duration_minutes=45,
        student_level="Basic",
        compliance_level="low",
        current_content="Content",
    )
    current_user = SimpleNamespace(id=456)
    before_files = set(get_settings().export_dir.glob("*.docx"))

    with pytest.raises(RuntimeError, match="record flush failed"):
        export_lesson_docx(FailingDb(), lesson=lesson, current_user=current_user)

    assert set(get_settings().export_dir.glob("*.docx")) == before_files
